from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def convert_md_to_docx(md_file, docx_file):
    """Convert markdown to docx with basic formatting"""
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    table = None
    for line in lines:
        line = line.rstrip()
        
        if not line:  # Empty line
            doc.add_paragraph()
            table = None # Reset table state
        elif line.strip().startswith('|') and '---' in line: # Table separator
            continue
        elif line.strip().startswith('|'): # Table row
            cells = [c.strip() for c in line.strip().split('|') if c]
            if not table:
                table = doc.add_table(rows=0, cols=len(cells))
                table.style = 'Table Grid'
            row_cells = table.add_row().cells
            for i, cell_text in enumerate(cells):
                if i < len(row_cells):
                    row_cells[i].text = cell_text
        elif line.startswith('# '):  # Heading 1
            p = doc.add_heading(line[2:], level=1)
            table = None
        elif line.startswith('## '):  # Heading 2
            p = doc.add_heading(line[3:], level=2)
            table = None
        elif line.startswith('### '):  # Heading 3
            p = doc.add_heading(line[4:], level=3)
            table = None
        elif line.startswith('#### '):  # Heading 4
            p = doc.add_heading(line[5:], level=4)
            table = None
        elif line.startswith('*   ') or line.startswith('- '):  # Bullet point
            p = doc.add_paragraph(line[4:] if line.startswith('*   ') else line[2:], style='List Bullet')
            table = None
        else:  # Regular paragraph
            p = doc.add_paragraph(line)
            table = None
    
    doc.save(docx_file)
    print(f"✅ Converted: {md_file} → {docx_file}")

# Convert all training files
files_to_convert = [
    '01_Science_and_Behavior.md',
    '02_Product_Catalog.md',
    '03_Consultation_and_Logistics.md',
    '04_Research_Library.md',
    '05_Protocol_Templates.md',
    '06_Provider_Directory_Structure.md',
    '07_Master_Research_Table.md',
    'exhaustive_faq.md',
    'PRODUCT_NAVIGATION_GUIDE.md',
    '10_Master_Protocol_Library.md'
]

for md_file in files_to_convert:
    docx_file = md_file.replace('.md', '_v2.docx')
    if os.path.exists(md_file):
        convert_md_to_docx(md_file, docx_file)
    else:
        print(f"⚠️ File not found: {md_file}")

print("\n✅ All files converted to .docx format!")
