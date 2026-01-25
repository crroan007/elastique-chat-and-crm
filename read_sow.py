import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
                
        return '\n'.join(full_text)
    except Exception as e:
        return str(e)

if __name__ == "__main__":
    file_path = r"C:\Homebrew Apps\Elastique - GPT_chatbot\Elastique - Chatbot SOW (2).docx"
    content = read_docx(file_path)
    with open("sow_content.txt", "w", encoding="utf-8") as f:
        f.write(content)
    print("SOW content saved to sow_content.txt")
